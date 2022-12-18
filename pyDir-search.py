import os
from docx import Document
from docx.shared import Pt

# Set the directory you want to search
root_dir = '/path/to/root/directory'

# Create a new Word document
document = Document()

# Set the font size and style for the document
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Function to create a bullet point in the document
def add_bullet_point(text):
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(20)
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(12)
    font.bold = False

# Function to create a heading in the document
def add_heading(text, level=1):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(14)
    font.bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = 1
    if level == 1:
        paragraph_format.space_before = Pt(20)
        paragraph_format.space_after = Pt(10)
    elif level == 2:
        paragraph_format.space_before = Pt(10)
        paragraph_format.space_after = Pt(10)

# Function to recursively search the directory tree and add the directory structure and file names to the document
def search_directory(root_dir):
    # Get the list of subdirectories and files in the current directory
    subdirs, files = [], []
    for name in os.listdir(root_dir):
        path = os.path.join(root_dir, name)
        if os.path.isdir(path):
            subdirs.append(name)
        else:
            files.append(name)

    # Add the current directory to the document
    add_heading(root_dir, level=2)

    # Add the list of files in the current directory to the document
    if len(files) > 0:
        add_heading('Files', level=3)
        for file in files:
            add_bullet_point(file)

    # Recursively search the subdirectories
    if len(subdirs) > 0:
        add_heading('Directories', level=3)
        for subdir in subdirs:
            search_directory(os.path.join(root_dir, subdir))

# Search the directory tree and add the results to the document
search_directory(root_dir)

# Save the document to the desktop
document.save('/path/to/desktop/all_Search.docx')
